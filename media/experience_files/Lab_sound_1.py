import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import sounddevice as sd
import soundfile as sf
import docx
from docx import Document
from docx.shared import Inches
from io import BytesIO

# data, fs = sf.read('sound1.wav', dtype='float32')
#
# print(data.dtype)
# print(data.shape)
#
# #zad1
# L=(data[:,0])
# R=(data[:,1])
# mix=(L+R)/2
#
# sf.write('sound_L.wav', L, fs)
# sf.write('sound_R.wav', R, fs)
# sf.write('sound_mix.wav', mix, fs)
#
# x=np.arange(0,data.shape[0])/fs
# plt.subplot(3,1,1)
# plt.plot(x,L)
#
# plt.subplot(3,1,2)
# plt.plot(x,R)
#
# plt.subplot(3,1,3)
# plt.plot(x,mix)
# plt.show()
#

# data, fs = sf.read('sin_440Hz.wav', dtype=np.int32)
# fsize=2**8
# plt.figure()
# plt.subplot(2,1,1)
# plt.plot(np.arange(0,data.shape[0])/fs,data)
# plt.subplot(2,1,2)
# yf = scipy.fftpack.fft(data,fsize)
# plt.plot(np.arange(0,fs/2,fs/fsize),20*np.log10( np.abs(yf[:fsize//2])))
# plt.show()

#zad2

data, fs = sf.read('sin_440Hz.wav', dtype=np.int32)
import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import soundfile as sf
from docx import Document
from docx.shared import Inches
from io import BytesIO

def plotAudio(Signal, Fs, fsize, axs, TimeMargin=[0, 0.02]):
    ax1 = axs[0]
    ax1.plot(np.arange(0, Signal.shape[0]) / Fs, Signal)
    ax1.set_xlabel("Czas (s)")
    ax1.set_ylabel("amplituda")
    ax1.set_title("Sygnał")
    ax1.set_xlim(TimeMargin[0], TimeMargin[1])

    yf = scipy.fftpack.fft(Signal, fsize)
    ax2 = axs[1]
    freqs = np.fft.fftfreq(fsize, 1/Fs)
    fft_vals = 20 * np.log10(np.clip(np.abs(yf[:fsize // 2]), a_min=1e-10, a_max=None))
    ax2.plot(np.arange(0, Fs / 2, Fs / fsize), 20 * np.log10(np.clip(np.abs(yf[:fsize // 2]), a_min=1e-10, a_max=None)))
    ax2.set_xlabel("Częstotliwość (Hz)")
    ax2.set_ylabel("decybele ")

    max_freq_idx = np.argmax(fft_vals)
    max_freq = max_freq_idx * (Fs / fsize)
    max_amp_val = fft_vals[max_freq_idx]

    return max_freq, max_amp_val

document = Document()
document.add_heading('Jan Szczudlo sprawozdanie', 0)

fsizes = [2**8, 2**12, 2**16]
files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
Margins = [[0, 0.02], [0.133, 0.155]]

for file in files:
    document.add_heading('Plik - {}'.format(file), 2)
    for fsize in fsizes:
        document.add_heading(f'Rozmiar fsize = {fsize}', 3)

        for Margin in Margins:
            data, fs = sf.read(file, dtype=np.int32)
            fig, axs = plt.subplots(2, 1, figsize=(10, 7))
            max_amp_freq, max_amp_val = plotAudio(data, fs, fsize, axs, Margin)



            fig.suptitle('Time margin {}'.format(Margin))  # Tytuł wykresu
            fig.tight_layout(pad=1.5)  # poprawa czytelności
            memfile = BytesIO()  # tworzenie bufora
            fig.savefig(memfile)  # z zapis do bufora

            document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

            memfile.close()

            document.add_paragraph(f'Najwyższa wartość widma: {max_amp_val} dB, Częstotliwość: {max_amp_freq} Hz')

document.save('report66.docx')

# Konwersja do formatu PDF
